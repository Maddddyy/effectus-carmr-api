"""
STAGE 0: Document Ingestion
Parse PDF, DOCX, TXT, and transcript files into clean text.
"""
import io
import os
from typing import List, Tuple
import logging

logger = logging.getLogger(__name__)


async def parse_document(filename: str, content: bytes) -> Tuple[str, str]:
    """
    Parse a document and return (clean_text, doc_type).
    doc_type: 'pdf' | 'docx' | 'txt' | 'transcript' | 'unknown'
    """
    ext = os.path.splitext(filename.lower())[1]

    if ext == ".pdf":
        return _parse_pdf(content), "pdf"
    elif ext in (".docx", ".doc"):
        return _parse_docx(content), "docx"
    elif ext in (".txt", ".md"):
        return content.decode("utf-8", errors="replace"), "txt"
    elif ext in (".vtt", ".srt"):
        return _parse_transcript(content), "transcript"
    else:
        # Try plain text fallback
        try:
            text = content.decode("utf-8", errors="replace")
            return text, "txt"
        except Exception:
            return "", "unknown"


def _parse_pdf(content: bytes) -> str:
    """Extract text from PDF using PyMuPDF."""
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=content, filetype="pdf")
        pages = []
        for page_num, page in enumerate(doc, 1):
            text = page.get_text()
            if text.strip():
                pages.append(f"[Page {page_num}]\n{text}")
        doc.close()
        return "\n\n".join(pages)
    except ImportError:
        logger.error("PyMuPDF not installed")
        return "[PDF parsing failed: PyMuPDF not available]"
    except Exception as e:
        logger.error(f"PDF parse error: {e}")
        return f"[PDF parsing failed: {e}]"


def _parse_docx(content: bytes) -> str:
    """Extract text from DOCX."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(content))
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        return "\n\n".join(paragraphs)
    except ImportError:
        logger.error("python-docx not installed")
        return "[DOCX parsing failed: python-docx not available]"
    except Exception as e:
        logger.error(f"DOCX parse error: {e}")
        return f"[DOCX parsing failed: {e}]"


def _parse_transcript(content: bytes) -> str:
    """Strip VTT/SRT timestamps and return clean transcript text."""
    import re
    text = content.decode("utf-8", errors="replace")
    # Remove VTT/SRT timestamps
    text = re.sub(r"\d{2}:\d{2}:\d{2}[.,]\d{3}\s*-->\s*\d{2}:\d{2}:\d{2}[.,]\d{3}", "", text)
    text = re.sub(r"^\d+$", "", text, flags=re.MULTILINE)
    text = re.sub(r"^WEBVTT.*$", "", text, flags=re.MULTILINE)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def truncate_for_context(text: str, max_chars: int = 80000) -> str:
    """
    Truncate document text to fit within LLM context.
    Preserves beginning and end (most important sections of board docs).
    """
    if len(text) <= max_chars:
        return text
    half = max_chars // 2
    return (
        text[:half]
        + f"\n\n[... {len(text) - max_chars} characters truncated for context ...]\n\n"
        + text[-half:]
    )


def combine_documents(texts: List[Tuple[str, str, str]]) -> str:
    """
    Combine multiple documents into a single context string.
    texts: List of (filename, clean_text, doc_type)
    """
    parts = []
    for filename, text, doc_type in texts:
        parts.append(f"=== DOCUMENT: {filename} (type: {doc_type}) ===\n\n{text}")
    return "\n\n" + ("=" * 60) + "\n\n".join(parts)
